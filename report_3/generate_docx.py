#!/usr/bin/env python3
"""
Generate a formatted DOCX from Integrated_Report_3.md
with yellow placeholder boxes for screenshots and ERDs.
"""

import re
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
MD_PATH = os.path.join(SCRIPT_DIR, "Integrated_Report_3.md")
OUT_PATH = os.path.join(SCRIPT_DIR, "Integrated_Report_3.docx")


def set_cell_shading(cell, color):
    """Set cell background color."""
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def add_placeholder_box(doc, text):
    """Add a yellow placeholder box for screenshots/ERDs."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x66, 0x33, 0x00)
    run.font.italic = True
    # Yellow background
    set_cell_shading(cell, "FFFFCC")
    # Set cell height
    tr = table.rows[0]._tr
    trPr = tr.get_or_add_trPr()
    trHeight = OxmlElement("w:trHeight")
    trHeight.set(qn("w:val"), "1200")
    trHeight.set(qn("w:hRule"), "atLeast")
    trPr.append(trHeight)
    doc.add_paragraph("")  # spacing


def add_code_block(doc, code_text):
    """Add a formatted code block."""
    p = doc.add_paragraph()
    p.style = doc.styles["Normal"]
    pf = p.paragraph_format
    pf.left_indent = Cm(1)
    pf.space_before = Pt(4)
    pf.space_after = Pt(4)
    # Set shading on paragraph
    pPr = p._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), "F5F5F5")
    shading.set(qn("w:val"), "clear")
    pPr.append(shading)
    run = p.add_run(code_text)
    run.font.name = "Consolas"
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)


def add_md_table(doc, header_line, separator_line, data_lines):
    """Add a markdown table as a Word table."""
    def parse_row(line):
        cells = [c.strip() for c in line.strip().strip("|").split("|")]
        return cells

    headers = parse_row(header_line)
    cols = len(headers)
    rows_data = [parse_row(l) for l in data_lines]

    table = doc.add_table(rows=1 + len(rows_data), cols=cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(9)
        set_cell_shading(cell, "D9E2F3")

    # Data rows
    for r_idx, row_data in enumerate(rows_data):
        for c_idx in range(min(cols, len(row_data))):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(row_data[c_idx])
            run.font.size = Pt(9)

    doc.add_paragraph("")  # spacing


def process_inline(paragraph, text):
    """Process inline markdown: **bold**, *italic*, `code`."""
    # Simple pattern matching for inline formatting
    parts = re.split(r'(\*\*.*?\*\*|\*.*?\*|`.*?`)', text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*") and not part.startswith("**"):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
        else:
            paragraph.add_run(part)


def main():
    with open(MD_PATH, "r", encoding="utf-8") as f:
        lines = f.readlines()

    doc = Document()

    # Page setup
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)

    # Default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)

    i = 0
    in_code_block = False
    code_lines = []
    in_tree_block = False
    tree_lines = []

    while i < len(lines):
        line = lines[i].rstrip("\n")

        # --- Code blocks ---
        if line.startswith("```"):
            if in_code_block:
                # End code block
                code_text = "\n".join(code_lines)
                add_code_block(doc, code_text)
                code_lines = []
                in_code_block = False
            elif line.strip() == "```" and i + 1 < len(lines) and (
                lines[i + 1].strip().startswith("SQL Server") or
                lines[i + 1].strip().startswith("├") or
                lines[i + 1].strip().startswith("│") or
                lines[i + 1].strip().startswith("└")
            ):
                in_code_block = True
            else:
                in_code_block = True
            i += 1
            continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        # --- Horizontal rule ---
        if line.strip() == "---":
            # Add a thin line
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), "6")
            bottom.set(qn("w:space"), "1")
            bottom.set(qn("w:color"), "999999")
            pBdr.append(bottom)
            pPr.append(pBdr)
            i += 1
            continue

        # --- Headings ---
        if line.startswith("# ") and not line.startswith("## "):
            p = doc.add_heading(line[2:].strip(), level=0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            i += 1
            continue
        if line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=1)
            i += 1
            continue
        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=2)
            i += 1
            continue
        if line.startswith("#### "):
            doc.add_heading(line[5:].strip(), level=3)
            i += 1
            continue

        # --- Screenshot / ERD placeholders ---
        placeholder_match = re.match(r'^\*\[(Screenshot.*?|INSERT:.*?)\]\*$', line.strip())
        if placeholder_match:
            add_placeholder_box(doc, placeholder_match.group(1))
            i += 1
            continue
        # Also catch lines like *[Screenshot 26:...]*
        placeholder_match2 = re.match(r'^\*\[(Screenshot.*?)\]\*$', line.strip())
        if placeholder_match2:
            add_placeholder_box(doc, placeholder_match2.group(1))
            i += 1
            continue

        # --- Tables ---
        if "|" in line and i + 1 < len(lines) and re.match(r'^\|[\s:-]+\|', lines[i + 1].strip()):
            header_line = line
            sep_line = lines[i + 1].rstrip("\n")
            data_lines = []
            j = i + 2
            while j < len(lines) and lines[j].strip().startswith("|"):
                data_lines.append(lines[j].rstrip("\n"))
                j += 1
            add_md_table(doc, header_line, sep_line, data_lines)
            i = j
            continue

        # --- Title metadata (bold lines at top) ---
        if line.startswith("**") and ":**" in line:
            p = doc.add_paragraph()
            process_inline(p, line)
            p.paragraph_format.space_after = Pt(2)
            i += 1
            continue

        # --- Bullet points ---
        if line.strip().startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            text = line.strip()[2:]
            process_inline(p, text)
            i += 1
            continue

        # --- Numbered lists ---
        num_match = re.match(r'^(\d+)\.\s+(.*)', line.strip())
        if num_match:
            p = doc.add_paragraph(style="List Number")
            process_inline(p, num_match.group(2))
            i += 1
            continue

        # --- Empty lines ---
        if line.strip() == "":
            i += 1
            continue

        # --- Regular paragraph ---
        p = doc.add_paragraph()
        process_inline(p, line)
        i += 1

    # Save
    doc.save(OUT_PATH)
    print(f"✅ DOCX saved to: {OUT_PATH}")
    print(f"   File size: {os.path.getsize(OUT_PATH) / 1024:.1f} KB")


if __name__ == "__main__":
    main()
