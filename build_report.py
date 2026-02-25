"""Build the full Consulting Report-1 Word document."""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
CHARTS = '/Users/bhavikdalal/Documents/data warehouse/project/report_charts'

def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x2F, 0x54, 0x96)

def add_body(text):
    doc.add_paragraph(text)

def add_image(filename, width=5.5):
    path = os.path.join(CHARTS, filename)
    if os.path.exists(path):
        doc.add_picture(path, width=Inches(width))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

def add_table(headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Light Grid Accent 1'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        t.rows[0].cells[i].text = h
        for p in t.rows[0].cells[i].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10)
    for row_data in rows:
        row = t.add_row()
        for i, val in enumerate(row_data):
            row.cells[i].text = str(val)
            for p in row.cells[i].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
    doc.add_paragraph()

def add_caption(text):
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        run.italic = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x59, 0x56, 0x59)

# ══════════════════════════════════════════════════════════
# COVER PAGE
# ══════════════════════════════════════════════════════════
for _ in range(6):
    doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('CONSULTING REPORT-1')
run.bold = True
run.font.size = Pt(24)
run.font.color.rgb = RGBColor(0x2F, 0x54, 0x96)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Requirements Gathering to Create Business Questions\nAnd Domain Understanding')
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0x2F, 0x54, 0x96)

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Design and Implementation of a Data Warehouse\nfor a Retail Store with Store-level Data')
run.font.size = Pt(13)

doc.add_paragraph()
for line in [
    'ISTM 637 — Data Warehousing',
    'Texas A&M University',
    'Spring 2026',
    '',
    'Team Members:',
    '[Member 1 Name]',
    '[Member 2 Name]',
    '[Member 3 Name]',
    '',
    'Date: February 25, 2026'
]:
    p = doc.add_paragraph(line)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_page_break()

# ══════════════════════════════════════════════════════════
# TABLE OF CONTENTS (placeholder)
# ══════════════════════════════════════════════════════════
add_heading('Table of Contents', level=1)
toc_items = [
    '1. Introduction .......................................................... 3',
    '2. Details about the Data ................................................ 4',
    '   2.1 Understanding of the Data ......................................... 4',
    '   2.2 Metadata for All OLTP Source Files ................................ 8',
    '   2.3 Entity Relationship Diagram ....................................... 12',
    '3. Retail Subject Area Understanding ..................................... 13',
    '   3.1 Literature Review ................................................. 13',
    '   3.2 Relevant Problems for DFF ......................................... 15',
    '4. Business Questions .................................................... 16',
    '   4.1 List of 10 Business Questions ..................................... 16',
    '   4.2 Data Evidence for Business Questions ............................... 20',
    '   4.3 Prioritization and Rationalization ................................. 24',
    'References ............................................................... 26',
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.style = doc.styles['Normal']
    for run in p.runs:
        run.font.name = 'Courier New'
        run.font.size = Pt(10)

doc.add_page_break()

# ══════════════════════════════════════════════════════════
# SECTION 1: INTRODUCTION
# ══════════════════════════════════════════════════════════
add_heading('1. Introduction', level=1)

add_body(
    "Dominick's Finer Foods (DFF) was a major supermarket chain operating approximately 100 stores "
    "in the Chicago metropolitan area. Between 1989 and 1994, the University of Chicago Booth School "
    "of Business (James M. Kilts Center for Marketing) partnered with Dominick's to conduct extensive "
    "store-level research on shelf management and pricing. Randomized experiments were conducted in "
    "over 25 product categories throughout the 100-store chain, producing one of the most comprehensive "
    "publicly available retail scanner datasets (Kilts Center, 2013, University of Chicago Booth)."
)

add_body(
    "The DFF dataset consists of approximately 5 GB of store-level data covering sales of more than "
    "3,500 UPCs (Universal Product Codes). The data includes weekly unit sales, retail prices, profit "
    "margins, deal/promotion codes, customer traffic counts, and store-level demographics derived from "
    "the 1990 U.S. Census. This dataset is notable for its comprehensive coverage and the inclusion "
    "of retail margin information, making it uniquely suitable for data warehouse design and OLAP analysis "
    "(https://www.chicagobooth.edu/research/kilts/datasets/dominicks)."
)

add_body(
    "A brief history of the company can be found at the Company Histories archive "
    "(https://www.company-histories.com/Dominicks-Finer-Foods-Inc-Company-History.html). "
    "Dominick's operated as a prominent regional grocery chain before being acquired by Safeway in 1998, "
    "and the stores were eventually closed in 2013."
)

add_body(
    "The objective of this consulting project is to design and develop a data warehouse for DFF that "
    "enables OLAP (Online Analytical Processing) analysis of its retail operations. This first report "
    "focuses on requirements gathering: understanding the data, researching the retail domain, and "
    "developing 10 business questions that the data warehouse will answer. These questions are designed "
    "to provide actionable insights for DFF management on pricing, promotions, category management, "
    "and store performance."
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════
# SECTION 2: DETAILS ABOUT THE DATA
# ══════════════════════════════════════════════════════════
add_heading('2. Details about the Data', level=1)

# ── 2.1 Understanding of the Data ──
add_heading('2.1 Understanding of the Data', level=2)

add_body(
    "To understand the DFF dataset, we imported sample data into Excel and performed exploratory "
    "analysis using Excel charts, Pivot Tables, and descriptive statistics. The following visualizations "
    "demonstrate key patterns in the data that inform our business questions."
)

# Chart 1
add_heading('Figure 1: Total Units Sold by Product Category', level=3)
add_image('chart1_category_volume.png')
add_caption('Figure 1: Total units sold by product category (sample: 200K rows per file)')
add_body(
    "Figure 1 shows the total units sold across 12 product categories sampled from the DFF dataset. "
    "Cigarettes and Canned Soup lead in volume from this sample, while Grooming Products shows the "
    "lowest volume. The wide variation in category volumes (from ~9K to over 8M units in sample) "
    "indicates that DFF categories have vastly different demand profiles, which must be accounted "
    "for in shelf space allocation and inventory management."
)

# Chart 2
add_heading('Figure 2: Weekly Sales Trend — Soft Drinks', level=3)
add_image('chart2_weekly_sdr_trend.png')
add_caption('Figure 2: Weekly total unit sales for Soft Drinks (SDR) category over time')
add_body(
    "Figure 2 displays the weekly unit sales trend for Soft Drinks, the largest category by row count "
    "(17.7 million rows). The data reveals significant weekly variation with periodic spikes that "
    "likely correspond to promotional events and seasonal demand patterns. These spikes — where "
    "weekly sales can reach over 150,000 units compared to a baseline of ~20,000 — demonstrate the "
    "dramatic impact of promotions on sales volume and validate the need for time-series OLAP analysis."
)

# Chart 3
add_heading('Figure 3: Promotion vs Non-Promotion Sales Comparison', level=3)
add_image('chart3_promo_vs_nonpromo.png')
add_caption('Figure 3: Average units sold during promoted vs non-promoted periods (Soft Drinks)')
add_body(
    "Figure 3 compares the average units sold when a promotion is active versus when no promotion "
    "is running. Promoted items sell an average of 133.60 units compared to 16.85 units for "
    "non-promoted items — a 7.9× sales lift. This confirms that promotions are highly effective "
    "at driving sales volume and validates the importance of promotion-related business questions."
)

# Chart 4
add_heading('Figure 4: Store Demographics — Urban/Suburban and Zone Distribution', level=3)
add_image('chart4_demographics.png', width=6.0)
add_caption('Figure 4: Distribution of DFF stores by urban/suburban classification and pricing zone')
add_body(
    "Figure 4 shows two key demographic dimensions of DFF's 107 stores. Approximately 38% of stores "
    "are urban and 62% are suburban. The pricing zone distribution shows Zone 2 has the highest "
    "concentration of stores, while several zones have only 1-3 stores. These demographic variations "
    "enable cross-store comparison analysis and support zone pricing optimization questions."
)

# Chart 5
add_heading('Figure 5: Descriptive Statistics by Category', level=3)
add_image('chart5_descriptive_stats.png', width=6.0)
add_caption('Figure 5: Descriptive statistics (mean, median, std dev, max, min, avg price, avg profit) by category')
add_body(
    "Figure 5 presents key descriptive statistics across six product categories. Notable findings include: "
    "(1) Soft Drinks have a very high standard deviation (117.27) relative to mean (17.07), indicating "
    "extreme sales variation driven by promotions; (2) Soft Drinks show negative average profit (-$8.28), "
    "suggesting loss-leader pricing strategy; (3) Cheeses have the highest average profit ($22.66), "
    "making them strategically important for margin management; and (4) Grooming Products have the "
    "lowest average movement (0.21), indicating slow-moving inventory."
)

# Chart 6
add_heading('Figure 6: Beer Price Spread Across Stores', level=3)
add_image('chart6_beer_price_spread.png')
add_caption('Figure 6: Top 10 Beer products by price spread (max - min unit price) across stores')
add_body(
    "Figure 6 identifies the Beer products with the largest price variation across DFF stores. "
    "Several products show price spreads exceeding $3-4, which may indicate pricing inconsistencies, "
    "zone pricing differences, or data quality issues. This type of pricing audit analysis is "
    "valuable for ensuring pricing compliance across the store chain."
)

doc.add_page_break()

# ── 2.2 Metadata ──
add_heading('2.2 Metadata for All OLTP Source Files', level=2)

add_body(
    "The DFF dataset consists of four types of OLTP source files organized into two categories: "
    "General Files (Customer Counts and Demographics) and Category-Specific Files (Movement and UPC). "
    "The following sections describe the metadata for each file type."
)

# Movement metadata
add_heading('A. Movement Files (Weekly Sales Data)', level=3)
add_body(
    "The Movement files contain weekly sales data for each UPC in each store, spanning over 5 years. "
    "There are 24 files — one for each product category — totaling approximately 134.9 million rows."
)

add_table(
    ['Column', 'Data Type', 'Description'],
    [
        ['UPC', 'Integer', 'Universal Product Code — join key to UPC files'],
        ['STORE', 'Integer', 'Store identifier — join key to Demographics'],
        ['WEEK', 'Integer', 'Proprietary week number (Week 1 ≈ Sep 14, 1989)'],
        ['MOVE', 'Integer', 'Units sold (movement) during the week'],
        ['QTY', 'Integer', 'Number of units for listed price (e.g., 2-for-$3 → QTY=2)'],
        ['PRICE', 'Decimal', 'Shelf price during the week'],
        ['SALE', 'Text', 'Deal flag: B=Bonus Buy, C=Coupon, S=Sale, blank=No promo'],
        ['PROFIT', 'Decimal', 'Gross profit for the week'],
        ['OK', 'Integer', 'Data quality flag (1 = valid observation)'],
    ]
)
add_caption('Table 1: Movement file column metadata')

add_body("Movement Files Inventory (24 files):")
movement_inventory = [
    ['Cookies (COO)', 'DONE-WCOO.csv', '13,447,807'],
    ['Soft Drinks (SDR)', 'wsdr.csv', '17,730,501'],
    ['Frozen Entrees (FRE)', 'WFRE-Done.csv', '11,663,663'],
    ['Cheeses (CHE)', 'Done-WCHE.csv', '9,427,395'],
    ['Analgesics (ANA)', 'DONE-WANA.csv', '7,339,217'],
    ['Canned Soup (CSO)', 'WCSO-Done.csv', '7,011,243'],
    ['Laundry Det. (LND)', 'wlnd.csv', '6,742,910'],
    ['Cereals (CER)', 'DONE-WCER.csv', '6,602,582'],
    ['Toothpaste (TPA)', 'WTPA_done.csv', '6,310,896'],
    ['Bottled Juices (BJC)', 'DONE-WBJC.csv', '6,222,797'],
    ['Cigarettes (CIG)', 'Done-WCIG.csv', '5,398,197'],
    ['Toothbrushes (TBR)', 'WTBR_done.csv', '4,529,484'],
    ['Fabric Softeners (FSF)', 'WFSF-done.csv', '4,122,487'],
    ['Beer (BER)', 'DONE-WBER.csv', '3,990,672'],
    ['Dish Detergent (DID)', 'WDID-Done.csv', '3,838,182'],
    ['Canned Tuna (TNA)', 'WTNA_done.csv', '3,763,229'],
    ['Crackers (CRA)', 'Done-WCRA.csv', '3,624,688'],
    ['Soaps (SOA)', 'WSOA.csv', '3,310,695'],
    ['Frozen Juices (FRJ)', 'WFRJ.csv', '3,156,545'],
    ['Grooming (GRO)', 'WGRO.csv', '2,016,291'],
    ['Bath Soap (BAT)', 'DONE-WBAT.csv', '1,644,557'],
    ['Toilet Tissue (TTI)', 'WTTI.csv', '1,627,284'],
    ['Front End Candies (FEC)', 'WFEC.csv', '797,921'],
    ['Frozen Dinners (FRD)', 'WFRD.csv', '562,715'],
]
add_table(['Category', 'Filename', 'Rows'], movement_inventory)
add_caption('Table 2: Movement files inventory — 24 files, ~134.9M total rows')

# UPC metadata
add_heading('B. UPC Files (Product Master Data)', level=3)
add_body(
    "The UPC files contain one record for each UPC in a category, with information about product "
    "name, size, commodity code, and case pack quantity. There are 28 files with approximately "
    "14,000 total UPCs."
)
add_table(
    ['Column', 'Data Type', 'Description'],
    [
        ['COM_CODE', 'Integer', 'Commodity/company code (brand/manufacturer grouping)'],
        ['UPC', 'Integer', 'Universal Product Code (join key to Movement files)'],
        ['DESCRIP', 'Text', 'Product description (truncated to ~20 characters)'],
        ['SIZE', 'Text', 'Package size (e.g., "7 OZ", "12 OZ", "1 CT")'],
        ['CASE', 'Integer', 'Case pack quantity'],
        ['NITEM', 'Integer', 'Internal DFF item number'],
    ]
)
add_caption('Table 3: UPC file column metadata')

# CCOUNT metadata
add_heading('C. Customer Count File (CCOUNT)', level=3)
add_body(
    "The Customer Count file contains store-level traffic data with 327,045 rows and 61 columns. "
    "Key columns include STORE, WEEK, CUSTCOUN (total customer count), and department-level counts "
    "(GROCERY, DAIRY, FROZEN, MEAT, PRODUCE, BAKERY, DELI, PHARMACY). The file also contains "
    "coupon redemption counts by department. Notable data quality issue: the DATE column is "
    "corrupted and contains garbled encoding, requiring the use of WEEK numbers for time analysis."
)

# DEMO metadata
add_heading('D. Demographics File (DEMO)', level=3)
add_body(
    "The Demographics file contains store-level attributes and census demographics with 108 rows "
    "(one per store) and 510 columns. Key columns include: STORE, NAME, CITY, ZIP (location); "
    "ZONE (pricing zone, 1-15), URBAN (binary flag), WEEKVOL (weekly volume); INCOME, EDUC, "
    "POVERTY, DENSITY (demographics); HSIZEAVG (household size); and PRICLOW/PRICMED/PRICHIGH "
    "(pricing tier flags). Note: 23 stores have null NAME and CITY values, and the first row "
    "uses dots (.) as placeholders for missing values."
)

# Relationships
add_heading('E. Key Relationships Among Data Files', level=3)
add_body("The OLTP source files are connected through the following join keys:")
add_table(
    ['Relationship', 'Join Key', 'Description'],
    [
        ['Movement → UPC', 'UPC', 'Links weekly sales to product details (name, size, brand)'],
        ['Movement → DEMO', 'STORE', 'Links sales data to store demographics and location'],
        ['Movement → CCOUNT', 'STORE + WEEK', 'Links sales to customer traffic for the same store-week'],
    ]
)
add_caption('Table 4: Join key relationships among OLTP source files')

# Data quality
add_heading('F. Data Quality Issues Identified', level=3)
add_table(
    ['Issue', 'File(s)', 'Severity', 'Impact'],
    [
        ['DATE column corrupted', 'CCOUNT', 'High', 'Cannot derive calendar dates; must use WEEK mapping'],
        ['Dots (.) as missing values', 'CCOUNT, DEMO', 'Medium', 'Columns read as text; requires cleaning'],
        ['SALE mostly blank (~90%)', 'Movement', 'Medium', 'Promotion analysis limited to ~10% of rows'],
        ['PRICE = 0 in some rows', 'Movement', 'Medium', 'Must filter; skews revenue calculations'],
        ['Non-UTF-8 encoding', 'UPC, CCOUNT', 'Low', 'Files require latin-1 or cp1252 encoding'],
        ['WEEK proprietary numbering', 'All', 'High', 'Week 1 ≈ Sep 14, 1989; requires date mapping formula'],
        ['23 null NAME/CITY values', 'DEMO', 'Low', 'ZIP and coordinates still available'],
    ]
)
add_caption('Table 5: Data quality issues identified across source files')

doc.add_page_break()

# ── 2.3 ERD ──
add_heading('2.3 Entity Relationship Diagram', level=2)
add_body(
    "The following ERD shows the relationships among the four OLTP source file types in the DFF "
    "dataset. Movement files are the central transactional entity, linked to UPC (product details), "
    "DEMO (store demographics), and CCOUNT (customer traffic) through the join keys identified above."
)
add_image('elaborate_erd.png', width=5.5)
add_caption('Figure 7: Entity Relationship Diagram — DFF OLTP Source Files')
add_body(
    "The Movement files serve as the primary fact data, recording weekly sales at the UPC-store-week "
    "grain. The UPC dimension provides product attributes (name, size, brand via COM_CODE). The DEMO "
    "dimension provides store-level demographics (zone, urban flag, income, density). The CCOUNT file "
    "provides customer traffic data that can be joined to Movement at the store-week level to analyze "
    "the relationship between foot traffic and sales performance."
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════
# SECTION 3: RETAIL SUBJECT AREA UNDERSTANDING
# ══════════════════════════════════════════════════════════
add_heading('3. Retail Subject Area Understanding', level=1)

add_heading('3.1 Literature Review', level=2)
add_body(
    "To understand the retail domain problems relevant to DFF, we reviewed published research "
    "that specifically used the Dominick's dataset or addressed comparable retail analytics challenges. "
    "The following five papers inform our business question development."
)

# Paper 1
add_heading('Pricing Strategy and Zone Pricing', level=3)
add_body(
    "Hoch, Kim, Montgomery, and Rossi (1995, Journal of Marketing Research, Vol. 32, No. 1, "
    "pp. 17-29) demonstrated that Dominick's implemented a micro-marketing pricing strategy, "
    "segmenting its stores into high, medium, and low price zones based on local competitive "
    "conditions. Their analysis showed that optimizing zone-level pricing based on category-specific "
    "demand elasticities could increase profits by 3-5% without increasing average prices. Our "
    "dataset includes 15 distinct pricing zones across 107 stores in the DEMO file, enabling us to "
    "validate and extend this analysis through business questions on cross-zone performance comparison."
)

# Paper 2
add_heading('Promotion Effectiveness and Sales Lifts', level=3)
add_body(
    "Chevalier, Kashyap, and Rossi (2003, American Economic Review, Vol. 93, No. 1, pp. 15-37) "
    "used the DFF dataset to study promotional pricing effects, finding that retail margins decrease "
    "during peak demand periods while sales volume increases substantially. Our statistical analysis "
    "confirms this: promotions boost Soft Drink sales by 7.9× and Canned Soup by 3.4× (see Section "
    "2.1, Figures 3 and 9). This validates questions on promotional lift measurement and customer "
    "traffic during promotions."
)

# Paper 3
add_heading('Consumer Demand Patterns and Stockpiling', level=3)
add_body(
    "Pesendorfer (2002, Journal of Business, Vol. 75, No. 1, pp. 33-66) analyzed cyclical pricing "
    "policies in the retail grocery sector and found that consumers strategically time purchases "
    "around promotional cycles. This stockpiling behavior explains why weekly sales MOVE values "
    "exhibit extreme variance (e.g., Soft Drinks: mean = 17.07, max = 4,637). Understanding these "
    "patterns is critical for DFF's inventory management and supports our time-series and trend "
    "analysis business questions."
)

# Paper 4
add_heading('Brand Competition and Market Share', level=3)
add_body(
    "Chintagunta (2002, Journal of Marketing Research, Vol. 39, No. 2, pp. 141-154) studied "
    "competitive brand interactions in retail grocery categories using scanner data and found that "
    "promotional activity by one brand significantly impacts competitors' market share within the "
    "same category. This supports our cross-brand market share analysis using the COM_CODE field "
    "in UPC files as a brand/manufacturer proxy."
)

# Paper 5
add_heading('Store Location, Demographics, and Performance', level=3)
add_body(
    "Montgomery (1997, Marketing Science, Vol. 16, No. 4, pp. 315-337) showed that store-level "
    "demographic characteristics — income, household size, education — are strong predictors of "
    "demand elasticity and category sales performance. The DFF DEMO file contains these exact "
    "variables for 107 stores, enabling demographic segmentation analysis. Our data confirms this: "
    "top-performing Toothpaste stores are 73.9% urban, suggesting density drives volume more than "
    "affluence."
)

# 3.2 Relevant Problems
add_heading('3.2 Relevant Retail Problems for DFF', level=2)
add_body("Based on the literature review, the following business challenges are most relevant for DFF:")

add_table(
    ['#', 'Problem', 'Relevance to DFF', 'Related BQs'],
    [
        ['1', 'Promotion ROI measurement', 'DFF runs Bonus Buys (B), Coupons (C), and Sales (S) across 24 categories', 'E3, M1'],
        ['2', 'Zone pricing optimization', 'DFF uses 15 pricing zones across 107 stores', 'H1, H3'],
        ['3', 'Category management', '24+ categories with vastly different profit profiles', 'E1, M3'],
        ['4', 'Trend detection & seasonality', '5+ years of weekly data enables time-series analysis', 'E2, H2'],
        ['5', 'Store performance benchmarking', 'Demographics + sales data enables performance attribution', 'M2, H1'],
    ]
)
add_caption('Table 6: Relevant retail problems for DFF with related business questions')

doc.add_page_break()

# ══════════════════════════════════════════════════════════
# SECTION 4: BUSINESS QUESTIONS
# ══════════════════════════════════════════════════════════
add_heading('4. Business Questions', level=1)

add_heading('4.1 List of 10 Business Questions', level=2)
add_body(
    "Based on our understanding of the data, metadata analysis, and domain research, we developed "
    "10 OLAP business questions organized by difficulty level. Each question is designed for "
    "OLAP-style aggregation, slicing, and dicing across the store, product, time, category, and "
    "promotion dimensions."
)

bqs = [
    ('E1', 'Easy', 'What were the total units sold per product category across the entire dataset period?',
     'Roll-up', 'All 24 Movement CSVs', 'MOVE, OK',
     'Category management foundation — determines shelf space allocation and supply chain priorities across all 24 categories.',
     "SELECT category_code, SUM(MOVE) AS total_units_sold\nFROM FactWeeklySales WHERE OK = 1\nGROUP BY category_code ORDER BY total_units_sold DESC;"),

    ('E2', 'Easy', 'What were the total weekly unit sales of Soft Drinks across all stores for each week?',
     'Roll-up', 'Movement/WSDR/wsdr.csv', 'WEEK, MOVE, OK',
     'High-volume trend monitoring — Soft Drinks is the largest category (17.7M rows). Weekly tracking enables demand forecasting and seasonal detection.',
     "SELECT WEEK, SUM(MOVE) AS total_units_sold\nFROM FactWeeklySales WHERE category_code = 'SDR' AND OK = 1\nGROUP BY WEEK ORDER BY WEEK;"),

    ('E3', 'Easy', 'How do promotion weeks compare to non-promotion weeks in terms of sales volume?',
     'Slice', 'Any Movement CSV (e.g., wsdr.csv)', 'SALE, MOVE, OK',
     'Promotion ROI baseline — data shows promotions boost SDR sales by 7.9×, confirming the investment value of promotional activities.',
     "SELECT CASE WHEN SALE IS NOT NULL THEN 'Promoted' ELSE 'Not Promoted' END AS promo_status,\n  AVG(MOVE) AS avg_units, SUM(MOVE) AS total_units\nFROM FactWeeklySales WHERE OK = 1\nGROUP BY promo_status;"),

    ('M1', 'Medium', 'Which promotion type (B=Bonus Buy, C=Coupon, S=Sale) generated the highest incremental unit sales lift in the Canned Soup category?',
     'Dice', 'Movement/WCSO/WCSO-Done.csv', 'UPC, STORE, WEEK, MOVE, SALE, OK',
     'Highest ROI impact — Sale discounts (S) drive 9.4× lift vs. baseline, 4.3× more effective than Bonus Buys. Directly informs promotional budget allocation. Supported by Chevalier et al. (2003).',
     "SELECT SALE AS promotion_type, AVG(MOVE) AS avg_units_promoted,\n  AVG(MOVE) - (SELECT AVG(MOVE) FROM FactWeeklySales\n    WHERE category_code='CSO' AND SALE IS NULL AND OK=1) AS sales_lift\nFROM FactWeeklySales\nWHERE category_code='CSO' AND SALE IS NOT NULL AND OK=1\nGROUP BY SALE ORDER BY sales_lift DESC;"),

    ('M2', 'Medium', 'How did quarterly revenue for Frozen Entrees differ between urban and suburban stores?',
     'Drill-down', 'WFRE/WFRE-Done.csv + Demographics/DEMO.csv', 'STORE, WEEK, MOVE, PRICE, QTY; URBAN',
     'Location strategy — suburban stores generate ~60% of revenue. Understanding differences informs store openings. Supported by Montgomery (1997).',
     "SELECT d.URBAN, (f.WEEK / 13) AS quarter_id,\n  SUM(f.MOVE * (f.PRICE / f.QTY)) AS quarterly_revenue\nFROM FactWeeklySales f JOIN DimStore d ON f.store_id = d.store_id\nWHERE f.category_code = 'FRE' AND f.OK = 1\nGROUP BY d.URBAN, quarter_id ORDER BY quarter_id;"),

    ('M3', 'Medium', 'How has the monthly market share (% of total units sold) of the top 5 Cereal brands changed over time?',
     'Pivot + Ratio-to-Total', 'DONE-WCER.csv + UPC/UPCCER.csv', 'UPC, WEEK, MOVE; COM_CODE',
     'Competitive intelligence — tracking brand share trends helps negotiate with manufacturers and optimize category assortment. Aligned with Chintagunta (2002).',
     "WITH brand_total AS (\n  SELECT u.COM_CODE, (f.WEEK/4) AS month_id, SUM(f.MOVE) AS brand_units\n  FROM FactWeeklySales f JOIN DimProduct u ON f.upc = u.upc\n  WHERE f.category_code='CER' AND f.OK=1\n  GROUP BY u.COM_CODE, month_id)\nSELECT COM_CODE, month_id, brand_units,\n  ROUND(100.0 * brand_units / SUM(brand_units) OVER(PARTITION BY month_id), 2) AS market_share_pct\nFROM brand_total;"),

    ('M4', 'Medium', 'Which 10 individual products (UPCs) in the Beer category had the largest price variance across stores in the same week?',
     'Dice + Ranking', 'DONE-WBER.csv + UPC/UPCBER.csv', 'UPC, STORE, WEEK, PRICE, QTY; DESCRIP',
     'Pricing consistency audit — Busch Beer shows $6.30 price spread across 69 stores, revealing genuine pricing inconsistency.',
     "SELECT f.UPC, p.DESCRIP, f.WEEK,\n  MAX(f.PRICE/f.QTY) - MIN(f.PRICE/f.QTY) AS price_spread\nFROM FactWeeklySales f JOIN DimProduct p ON f.upc = p.upc\nWHERE f.category_code='BER' AND f.OK=1 AND f.PRICE>0\nGROUP BY f.UPC, p.DESCRIP, f.WEEK\nHAVING COUNT(DISTINCT f.STORE) >= 5\nORDER BY price_spread DESC LIMIT 10;"),

    ('H1', 'Hard', 'Which stores fall into the top 25%, middle 50%, and bottom 25% of total revenue for the Toothpaste category, and how do their demographics differ?',
     'NTILE + Drill-down', 'WTPA/WTPA_done.csv + Demographics/DEMO.csv', 'STORE, MOVE, PRICE, QTY; INCOME, URBAN, ZONE',
     'Zone pricing validation — top-performing stores are 73.9% urban. Paradoxically, lower income areas drive higher volume. Validates Hoch et al. (1995).',
     "WITH store_revenue AS (\n  SELECT STORE, SUM(MOVE*(PRICE/QTY)) AS total_revenue\n  FROM FactWeeklySales WHERE category_code='TPA' AND OK=1 AND PRICE>0\n  GROUP BY STORE),\nquartiles AS (\n  SELECT STORE, total_revenue,\n    NTILE(4) OVER(ORDER BY total_revenue) AS revenue_quartile\n  FROM store_revenue)\nSELECT CASE WHEN revenue_quartile=1 THEN 'Bottom 25%'\n  WHEN revenue_quartile IN(2,3) THEN 'Middle 50%'\n  ELSE 'Top 25%' END AS tier,\n  COUNT(*) AS store_count, AVG(total_revenue) AS avg_revenue\nFROM quartiles GROUP BY tier;"),

    ('H2', 'Hard', 'For each week, rank the top 10 individual products (UPCs) in the Crackers category by unit sales, and show their week-over-week sales change.',
     'RANK + LAG', 'WCRA/Done-WCRA.csv + UPC/UPCCRA.csv', 'UPC, WEEK, MOVE; DESCRIP',
     'Product velocity monitoring — identifies rising/falling products weekly, enabling rapid assortment adjustment and measuring promotional impact on individual SKUs.',
     "WITH weekly_product AS (\n  SELECT f.UPC, p.DESCRIP, f.WEEK, SUM(f.MOVE) AS weekly_units,\n    RANK() OVER(PARTITION BY f.WEEK ORDER BY SUM(f.MOVE) DESC) AS week_rank\n  FROM FactWeeklySales f JOIN DimProduct p ON f.upc=p.upc\n  WHERE f.category_code='CRA' AND f.OK=1\n  GROUP BY f.UPC, p.DESCRIP, f.WEEK)\nSELECT *, LAG(weekly_units) OVER(PARTITION BY UPC ORDER BY WEEK) AS prev_week,\n  weekly_units - LAG(weekly_units) OVER(PARTITION BY UPC ORDER BY WEEK) AS wow_change\nFROM weekly_product WHERE week_rank <= 10;"),

    ('H3', 'Hard', 'For each pricing zone, what is the percentile rank of each store\'s average weekly profit in the Cigarettes category, and how does it compare to the zone average?',
     'PERCENT_RANK + Drill-down', 'WCIG/Done-WCIG.csv + Demographics/DEMO.csv', 'STORE, WEEK, PROFIT; ZONE',
     'Profit optimization — Cigarettes high-margin (avg $16.44). Zone 1 has widest profit spread ($7–$40), showing significant underperformance by some stores.',
     "WITH store_avg AS (\n  SELECT f.STORE, d.ZONE, AVG(f.PROFIT) AS avg_weekly_profit\n  FROM FactWeeklySales f JOIN DimStore d ON f.store_id=d.store_id\n  WHERE f.category_code='CIG' AND f.OK=1\n  GROUP BY f.STORE, d.ZONE)\nSELECT STORE, ZONE, avg_weekly_profit,\n  PERCENT_RANK() OVER(PARTITION BY ZONE ORDER BY avg_weekly_profit) AS pctile_rank,\n  AVG(avg_weekly_profit) OVER(PARTITION BY ZONE) AS zone_avg\nFROM store_avg ORDER BY ZONE, pctile_rank DESC;"),
]

for bq_id, diff, question, olap, files, cols, value, sql in bqs:
    add_heading(f'BQ {bq_id}: {question}', level=3)
    add_table(
        ['Item', 'Detail'],
        [
            ['Difficulty', diff],
            ['OLAP Operation', olap],
            ['Files Needed', files],
            ['Key Columns', cols],
            ['Business Value', value],
        ]
    )
    p = doc.add_paragraph()
    run = p.add_run('SQL Query:')
    run.bold = True
    p = doc.add_paragraph()
    run = p.add_run(sql)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)

doc.add_page_break()

# ── 4.2 Data Evidence ──
add_heading('4.2 Data Evidence for Business Questions', level=2)
add_body(
    "The following charts, generated from the actual DFF dataset, provide data-backed evidence "
    "that each business question is implementable and produces meaningful results."
)

evidence = [
    ('chart1_category_volume.png', 'BQ E1: Category Volume Comparison',
     'Figure 8: Total units sold by category confirms wide volume variation, validating the category management question.',
     'This chart directly supports BQ E1. The data shows a 100× difference between the highest and lowest volume categories, confirming that category-level analysis is essential for DFF.'),
    ('chart2_weekly_sdr_trend.png', 'BQ E2: Weekly Sales Trend',
     'Figure 9: Weekly Soft Drinks sales trend with periodic promotional spikes.',
     'This chart supports BQ E2 by showing clear weekly patterns and promotional spikes in the SDR data.'),
    ('chart3_promo_vs_nonpromo.png', 'BQ E3: Promotion vs Non-Promotion',
     'Figure 10: Average units sold by promotion status — 7.9× lift for promoted items.',
     'This directly validates BQ E3. Promoted items sell 7.9× more units on average, confirming promotion effectiveness.'),
    ('bq_m1_promo_lift.png', 'BQ M1: Promotion Lift by Deal Type',
     'Figure 11: Average units sold by deal type for Canned Soup.',
     'This chart supports BQ M1 by showing that different deal types (B, S, C) produce significantly different sales lifts, with Sale (S) discounts being the most effective.'),
    ('bq_m2_urban_suburban.png', 'BQ M2: Urban vs Suburban Revenue',
     'Figure 12: Quarterly revenue comparison between urban and suburban stores (Frozen Entrees).',
     'This dual-bar chart supports BQ M2 by showing consistent revenue differences between urban and suburban locations across quarters.'),
    ('bq_h1_quartiles.png', 'BQ H1: Store Revenue Quartiles',
     'Figure 13: Store quartile analysis with demographic overlay (Toothpaste).',
     'This supports BQ H1 by showing that top-performing stores have dramatically different demographic profiles than bottom performers, validating the NTILE segmentation approach.'),
    ('bq_h3_zone_profit.png', 'BQ H3: Profit by Pricing Zone',
     'Figure 14: Average store profit by pricing zone for Cigarettes with min/max range.',
     'This supports BQ H3 by showing significant profit variation both between and within zones, validating the percentile ranking approach.'),
]

for img, title, caption, interp in evidence:
    add_heading(title, level=3)
    add_image(img)
    add_caption(caption)
    add_body(interp)

doc.add_page_break()

# ── 4.3 Prioritization ──
add_heading('4.3 Prioritization and Rationalization', level=2)
add_body(
    "The 10 business questions are prioritized below based on strategic value to DFF management, "
    "data availability, and implementability."
)

add_table(
    ['Priority', 'BQ', 'Rationale'],
    [
        ['1', 'M1 — Promo lift by deal type', 'Highest ROI — directly impacts promotional budget. Data confirms 9.4× lift for Sale vs 2.2× for Bonus Buy. Supported by Chevalier et al. (2003).'],
        ['2', 'E1 — Total units by category', 'Category management foundation: baseline for shelf space, supply chain, and merchandising across 24 categories.'],
        ['3', 'H1 — Store quartiles + demographics', 'Zone pricing validation: Hoch et al. (1995) showed 3-5% profit improvement. Top stores are 73.9% urban.'],
        ['4', 'M3 — Brand market share', 'Competitive intelligence for manufacturer negotiation. Aligned with Chintagunta (2002).'],
        ['5', 'E2 — Weekly SDR trend', 'Largest category — demand forecasting and inventory optimization for 17.7M-row dataset.'],
        ['6', 'M2 — Urban vs suburban revenue', 'Location strategy: 60/40 suburban-urban revenue split informs store openings. Supported by Montgomery (1997).'],
        ['7', 'H3 — Percentile rank by zone', 'Profit optimization: Cigarettes high-margin ($16.44 avg). Zone 1 spread $7-$40 shows improvement opportunity.'],
        ['8', 'E3 — Promo vs non-promo', 'Validates promotional ROI: 7.9× lift proves promotions worth the investment.'],
        ['9', 'H2 — Top 10 products WoW', 'Product velocity: enables rapid assortment adjustment based on weekly movement trends.'],
        ['10', 'M4 — Price variance (Beer)', 'Pricing audit: $6.30 spread across 69 stores reveals pricing inconsistency requiring correction.'],
    ]
)
add_caption('Table 7: Business question prioritization with strategic rationale')

add_body(
    "Prioritization Rationale: Our prioritization follows four principles: (1) Revenue and profit "
    "impact first — questions that directly inform pricing, promotion, and category management decisions "
    "rank highest because they impact DFF's top and bottom line. (2) Strategic alignment — questions "
    "supported by published research on the DFF dataset (Hoch et al., 1995; Chevalier et al., 2003) "
    "receive higher priority because they address validated business problems. (3) Data richness — "
    "questions leveraging the most complete and granular data (Movement files with 134.9M rows) are "
    "ranked higher than those dependent on files with quality issues. (4) Actionability — questions "
    "that produce immediately actionable insights are prioritized over descriptive questions."
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════
# REFERENCES
# ══════════════════════════════════════════════════════════
add_heading('References', level=1)

refs = [
    'Chevalier, J.A., Kashyap, A.K., & Rossi, P.E. (2003). "Why Don\'t Prices Rise During Periods of Peak Demand? Evidence from Scanner Data." American Economic Review, Vol. 93, No. 1, pp. 15-37.',
    'Chintagunta, P.K. (2002). "Investigating Category Pricing Behavior at a Retail Chain." Journal of Marketing Research, Vol. 39, No. 2, pp. 141-154.',
    'Hoch, S.J., Kim, B.D., Montgomery, A.L., & Rossi, P.E. (1995). "Determinants of Store-Level Price Elasticity." Journal of Marketing Research, Vol. 32, No. 1, pp. 17-29.',
    'Montgomery, A.L. (1997). "Creating Micro-Marketing Pricing Strategies Using Supermarket Scanner Data." Marketing Science, Vol. 16, No. 4, pp. 315-337.',
    'Pesendorfer, M. (2002). "Retail Sales: A Study of Pricing Behavior in Supermarkets." Journal of Business, Vol. 75, No. 1, pp. 33-66.',
    'Kilts Center for Marketing (2013). Dominick\'s Data Manual and Codebook. University of Chicago Booth School of Business. Available at: https://www.chicagobooth.edu/research/kilts/datasets/dominicks',
    'Dominick\'s Finer Foods Inc. Company History. Available at: https://www.company-histories.com/Dominicks-Finer-Foods-Inc-Company-History.html',
]

for i, ref in enumerate(refs, 1):
    add_body(f"{i}. {ref}")

# ══════════════════════════════════════════════════════════
# SAVE
# ══════════════════════════════════════════════════════════
output_path = '/Users/bhavikdalal/Documents/data warehouse/project/Consulting_Report_1_DFF.docx'
doc.save(output_path)
print(f"\n✅ REPORT SAVED: {output_path}")
print(f"   Sections: Introduction, Data Details (Understanding + Metadata + ERD), Subject Area, Business Questions")
print(f"   Charts embedded: 11")
print(f"   Tables: 7")
print(f"   References: {len(refs)}")
